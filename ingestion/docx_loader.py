"""
DOCX loader with hierarchical structure and image extraction
"""

from typing import List, Dict, Tuple
from pathlib import Path
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os
from utils.logging_config import get_logger

logger = get_logger(__name__)


class DOCXLoader:
    """
    Load DOCX files with:
    - Hierarchical structure (H1, H2, H3)
    - Image extraction
    - Text content
    """
    
    def __init__(self, output_image_dir: str = "data/images"):
        self.output_image_dir = Path(output_image_dir)
        self.output_image_dir.mkdir(parents=True, exist_ok=True)
    
    def load(self, docx_path: str) -> Dict:
        """
        Load DOCX file and extract structured content
        
        Returns:
            Dict with:
            - sections: Hierarchical structure
            - images: List of extracted images
            - metadata: Document metadata
        """
        logger.info(f"Loading DOCX: {docx_path}")
        
        doc = Document(docx_path)
        doc_name = Path(docx_path).stem
        
        # Extract images
        images = self._extract_images(doc, doc_name)
        logger.info(f"Extracted {len(images)} images")
        
        # Extract hierarchical structure
        sections = self._extract_sections(doc, images)
        logger.info(f"Extracted {len(sections)} top-level sections")
        
        # Extract metadata
        metadata = {
            'filename': Path(docx_path).name,
            'total_paragraphs': len(doc.paragraphs),
            'total_images': len(images),
            'total_sections': len(sections)
        }
        
        return {
            'sections': sections,
            'images': images,
            'metadata': metadata
        }
    
    def _extract_images(self, doc: Document, doc_name: str) -> List[Dict]:
        """Extract images from DOCX"""
        images = []
        image_counter = 0
        
        # Get image relationships
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    
                    # Determine file extension
                    content_type = rel.target_part.content_type
                    ext = {
                        'image/png': '.png',
                        'image/jpeg': '.jpg',
                        'image/jpg': '.jpg',
                        'image/gif': '.gif',
                        'image/bmp': '.bmp'
                    }.get(content_type, '.png')
                    
                    # Save image
                    image_filename = f"{doc_name}_image_{image_counter}{ext}"
                    image_path = self.output_image_dir / image_filename
                    
                    with open(image_path, 'wb') as f:
                        f.write(image_data)
                    
                    images.append({
                        'id': f"img_{image_counter}",
                        'filename': image_filename,
                        'path': str(image_path),
                        'index': image_counter
                    })
                    
                    image_counter += 1
                
                except Exception as e:
                    logger.warning(f"Failed to extract image: {e}")
        
        return images
    
    def _extract_sections(self, doc: Document, images: List[Dict]) -> List[Dict]:
        """Extract hierarchical section structure"""
        sections = []
        current_h1 = None
        current_h2 = None
        current_h3 = None
        image_index = 0
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            style = para.style.name
            
            # Heading 1
            if style == 'Heading 1':
                current_h1 = {
                    'level': 1,
                    'title': text,
                    'text': '',
                    'subsections': [],
                    'images': [],
                    'para_index': para_idx
                }
                sections.append(current_h1)
                current_h2 = None
                current_h3 = None
            
            # Heading 2
            elif style == 'Heading 2' and current_h1:
                current_h2 = {
                    'level': 2,
                    'title': text,
                    'text': '',
                    'subsections': [],
                    'images': [],
                    'para_index': para_idx
                }
                current_h1['subsections'].append(current_h2)
                current_h3 = None
            
            # Heading 3
            elif style == 'Heading 3' and current_h2:
                current_h3 = {
                    'level': 3,
                    'title': text,
                    'text': '',
                    'subsections': [],
                    'images': [],
                    'para_index': para_idx
                }
                current_h2['subsections'].append(current_h3)
            
            # Normal text
            elif style in ['Normal', 'Body Text', 'normal']:
                # Add to most specific current section
                if current_h3:
                    current_h3['text'] += text + '\n\n'
                elif current_h2:
                    current_h2['text'] += text + '\n\n'
                elif current_h1:
                    current_h1['text'] += text + '\n\n'
        
        # Assign images to sections (simple strategy: by proximity)
        self._assign_images_to_sections(sections, images, len(doc.paragraphs))
        
        return sections
    
    def _assign_images_to_sections(self, sections: List[Dict], images: List[Dict], total_paras: int):
        """Assign images to sections based on position"""
        # Simple strategy: evenly distribute images across sections
        if not images:
            return
        
        images_per_section = len(images) / max(len(sections), 1)
        
        for i, section in enumerate(sections):
            start_img = int(i * images_per_section)
            end_img = int((i + 1) * images_per_section)
            section['images'] = images[start_img:end_img]
            
            # Recursively assign to subsections
            if section['subsections']:
                subsection_images = section['images']
                imgs_per_subsection = len(subsection_images) / max(len(section['subsections']), 1)
                
                for j, subsection in enumerate(section['subsections']):
                    start = int(j * imgs_per_subsection)
                    end = int((j + 1) * imgs_per_subsection)
                    subsection['images'] = subsection_images[start:end]